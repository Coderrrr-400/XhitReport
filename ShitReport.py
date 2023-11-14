import sys
import tempfile
from PyQt5.QtWidgets import QApplication, QWidget, QLabel, QLineEdit, QComboBox, QPushButton, QVBoxLayout, QHBoxLayout, QFormLayout, QMessageBox
from PyQt5.QtGui import QPixmap, QIcon
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from PIL import Image

class ReportGenerator(QWidget):
    # 定义样式的类变量
    HEADING_STYLE = 'MyHeading'
    NORMAL_STYLE = 'Normal'
    NORMAL_2_STYLE = 'Normal_2'

    def __init__(self):
        super().__init__()
        self.setWindowIcon(QIcon('th.jpg'))
        self.init_ui()

    def init_ui(self):
        self.labels = ["漏洞所属:", "漏洞类型:", "目标系统:", "漏洞URL:", "IP地址:"]
        self.text_edits = [QLineEdit(self) for _ in range(5)]
        self.combo_box = QComboBox(self)
        self.combo_box.addItems(['文件上传', '远程命令执行', 'SQL注入', '弱口令', '敏感信息泄露', '任意文件下载', '任意文件读取', '反序列化', 'SSRF'])
        self.image_label_asset = QLabel(self)
        self.paste_button_asset = QPushButton('点击读取截图', self)
        self.paste_button_asset.clicked.connect(lambda: self.paste_image('asset'))
        self.image_label_vuln = QLabel(self)
        self.paste_button_vuln = QPushButton('点击读取截图', self)
        self.paste_button_vuln.clicked.connect(lambda: self.paste_image('vuln'))
        self.text_edit_vuln_2 = QLineEdit(self)
        self.image_label_vuln_2 = QLabel(self)
        self.paste_button_vuln_2 = QPushButton('点击读取截图', self)
        self.paste_button_vuln_2.clicked.connect(lambda: self.paste_image('vuln_2'))
        self.repair_suggestion_text_edit = QLineEdit(self)

        self.generate_button = QPushButton('生成报告', self)
        self.generate_button.clicked.connect(self.generate_report)

        form_layout = QFormLayout()
        form_layout.addRow(QLabel(self.labels[0]), self.text_edits[0])
        form_layout.addRow(QLabel(self.labels[1]), self.combo_box)

        type_system_layout = QHBoxLayout()
        type_system_layout.addWidget(QLabel(self.labels[2]))
        type_system_layout.addWidget(self.text_edits[1])
        type_system_layout.addWidget(QLabel(self.labels[3]))
        type_system_layout.addWidget(self.text_edits[2])
        form_layout.addRow(type_system_layout)

        form_layout.addRow(QLabel(self.labels[4]), self.text_edits[3])

        form_layout.addRow(QLabel("资产证明图片:"), self.image_label_asset)
        form_layout.addRow(self.paste_button_asset)
        form_layout.addRow(QLabel("漏洞复现描述:"), self.text_edits[4])
        form_layout.addRow(QLabel("漏洞证明图片:"), self.image_label_vuln)
        form_layout.addRow(self.paste_button_vuln)
        form_layout.addRow(QLabel("漏洞复现描述2:"), self.text_edit_vuln_2)
        form_layout.addRow(QLabel("漏洞证明图片2:"), self.image_label_vuln_2)
        form_layout.addRow(self.paste_button_vuln_2)
        form_layout.addRow(QLabel("漏洞修复建议:"), self.repair_suggestion_text_edit)

        button_layout = QHBoxLayout()
        button_layout.addStretch(1)
        button_layout.addWidget(self.generate_button)

        main_layout = QVBoxLayout()
        main_layout.addLayout(form_layout)
        main_layout.addLayout(button_layout)

        self.setLayout(main_layout)
        self.setWindowTitle('ShitReport by:Size0f')
        self.setFixedSize(800, 800)

        self.paste_asset = False
        self.paste_vuln = False
        self.paste_vuln_2 = False

        clipboard = QApplication.clipboard()
        clipboard.dataChanged.connect(self.on_clipboard_changed)

    def on_clipboard_changed(self):
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            pixmap = QPixmap(mime_data.imageData())
            scaled_pixmap = pixmap.scaled(400, 300, aspectRatioMode=Qt.AspectRatioMode.KeepAspectRatio)

            if self.paste_asset:
                self.image_label_asset.setPixmap(scaled_pixmap)
                self.original_pixmap_asset = pixmap
                self.paste_asset = False
            elif self.paste_vuln:
                self.image_label_vuln.setPixmap(scaled_pixmap)
                self.original_pixmap_vuln = pixmap
                self.paste_vuln = False
            elif self.paste_vuln_2:
                self.image_label_vuln_2.setPixmap(scaled_pixmap)
                self.original_pixmap_vuln_2 = pixmap
                self.paste_vuln_2 = False

    def paste_image(self, image_type):
        if image_type == 'asset':
            self.paste_asset = True
        elif image_type == 'vuln':
            self.paste_vuln = True
        elif image_type == 'vuln_2':
            self.paste_vuln_2 = True
            
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            pixmap = QPixmap(mime_data.imageData())
            scaled_pixmap = pixmap.scaled(400, 300, aspectRatioMode=Qt.AspectRatioMode.KeepAspectRatio)

            if self.paste_asset:
                self.image_label_asset.setPixmap(scaled_pixmap)
                self.original_pixmap_asset = pixmap
                self.paste_asset = False
            elif self.paste_vuln:
                self.image_label_vuln.setPixmap(scaled_pixmap)
                self.original_pixmap_vuln = pixmap
                self.paste_vuln = False
            elif self.paste_vuln_2:
                self.image_label_vuln_2.setPixmap(scaled_pixmap)
                self.original_pixmap_vuln_2 = pixmap
                self.paste_vuln_2 = False

    def generate_report(self):
        vuln_belongs, target_system, vuln_url, ip_address = [text_edit.text() for text_edit in self.text_edits[:-1]]
        vuln_type = self.combo_box.currentText()
        pixmap_asset = self.original_pixmap_asset if hasattr(self, 'original_pixmap_asset') else None
        vuln_proof = self.text_edits[-1].text()
        pixmap_vuln = self.original_pixmap_vuln if hasattr(self, 'original_pixmap_vuln') else None
        vuln_proof_2 = self.text_edit_vuln_2.text()
        pixmap_vuln_2 = self.original_pixmap_vuln_2 if hasattr(self, 'original_pixmap_vuln_2') else None
        repair_suggestion = self.repair_suggestion_text_edit.text()
        

        image_path_asset, image_path_vuln, image_path_vuln_2 = '', '', ''
        if pixmap_asset:
            temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            image_path_asset = temp_file.name
            pixmap_asset.save(image_path_asset)

        if pixmap_vuln:
            temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            image_path_vuln = temp_file.name
            pixmap_vuln.save(image_path_vuln)
        
        if pixmap_vuln_2:
            temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            image_path_vuln_2 = temp_file.name
            pixmap_vuln_2.save(image_path_vuln_2)

        vuln_descriptions = {
            '文件上传': '任意文件上传是指：攻击者利用该漏洞在受影响的网站上上传恶意文件或包含恶意代码的文件。这些文件可能被执行或访问，导致攻击者在服务器上执行任意命令、获取敏感信息或控制整个网站。漏洞通常由于上传功能未经充分验证和过滤而产生，使攻击者可以绕过文件类型和大小的限制。',
            '远程命令执行': '远程命令执行是由于开发人员编写源码，没有针对代码中可执行的特殊函数入口做过滤，导致客户端可以提交恶意构造语句提交，并交由服务器端执行。命令注入攻击中WEB服务器没有过滤类似system(),eval()，exec()等函数是该漏洞攻击成功的最主要原因。',
            'SQL注入': '结构化查询语言（Structured Query Language，简称SQL）是一种特殊目的的编程语言，是一种数据库查询和程序设计语言，用于存取数据以及查询、更新和管理关系型数据库。SQL注入漏洞主要形成的原因是Web应用程序对用户的输入没有做严格的判断，导致用户可用将非法的SQL语句拼接到正常的语句中，被当作SQL语句的一部分执行。',
            '弱口令': '弱⼝令即容易破译的密码，多为简单的数字组合、帐号相同的数字组合、键盘上的临近键或常见姓名、终端设备出⼚配置通⽤密码等都属于弱密码范畴。攻击者很容易预测⽤户名和密码，登录应⽤程序，从⽽获取未获授权的特权。',
            '敏感信息泄露': '敏感信息泄露漏洞是指攻击者能够非法地访问和获取系统中的敏感信息，这些信息可能包括但不限于用户凭证、个人身份信息、财务数据等。这种漏洞的发现可能导致用户隐私的泄露，进而可能对个人或组织造成严重损害。攻击者通常通过利用应用程序或系统中存在的安全漏洞，获取未经授权的访问权限，从而访问到敏感信息。',
            '任意文件下载': '任意文件下载漏洞允许攻击者通过应用程序获取服务器上的任何文件，包括敏感数据、配置文件和源代码。攻击者可以借此窃取敏感信息，执行恶意代码，或进一步渗透系统。危害包括数据泄露、系统瘫痪，降低可用性，以及潜在的后门安装，损害机密性、完整性和可用性，可能导致法律责任和声誉损害。修复建议包括实施强制访问控制，限制文件下载路径，以及对输入进行适当验证和过滤。',
            '任意文件读取': '任意文件读取漏洞是指攻击者通过未正确验证用户输入，利用漏洞读取服务器上的任意文件。攻击者可能通过遍历目录或使用特殊构造的路径来获取敏感信息，如配置文件、密码文件等。这漏洞可能导致泄露敏感数据，进而危及系统的安全性和隐私。',
            '反序列化': '反序列化漏洞是基于序列化和反序列化的操作，在反序列化——unserialize()时存在用户可控参数，而反序列化会自动调用一些魔术方法，如果魔术方法内存在一些敏感操作例如eval()函数，而且参数是通过反序列化产生的，那么用户就可以通过改变参数来执行敏感操作，这就是反序列化漏洞。',
            'SSRF': 'SSRF (Server-Side Request Forgery，服务器端请求伪造) 是一种由攻击者构造请求，由服务端发起请求的安全漏洞，一般情况下，SSRF攻击的目标是外网无法访问的内网系统，也正因为请求是由服务端发起的，所以服务端能请求到与自身相连而与外网隔绝的内部系统。也就是说可以利用一个网络请求的服务，当作跳板进行攻击。'

        }

        self.add_image_to_document(image_path_asset, image_path_vuln, image_path_vuln_2, vuln_belongs, vuln_type, target_system, vuln_url, ip_address, vuln_proof, vuln_proof_2, repair_suggestion, vuln_descriptions)

    def add_image_to_document(self, image_path_asset, image_path_vuln, image_path_vuln_2, vuln_belongs, vuln_type, target_system, vuln_url, ip_address, vuln_proof, vuln_proof_2, repair_suggestion, vuln_descriptions):
        try:
            doc = Document()

            # Define a new style named 'MyHeading' for Heading 1
            my_heading_style = doc.styles.add_style('MyHeading', 1)
            my_heading_style.font.name = '等线'
            my_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            my_heading_style.font.size = Pt(16)
            my_heading_style.font.bold = True
            my_heading_style.paragraph_format.alignment = 1

            # Define and apply 'Normal' style
            normal_style = doc.styles['Normal']
            normal_style.font.name = u'宋体'
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            normal_style.font.size = Pt(10.5)

            # Define and apply 'Normal_2' style (if needed)
            normal_2_style = doc.styles.add_style('Normal_2', 1)  # Add a new style named 'Normal_2'
            normal_2_style.font.name = u'宋体'
            normal_2_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            normal_2_style.font.size = Pt(14)

            # Apply font styles to paragraphs
            doc.add_paragraph(f'{vuln_belongs}{vuln_type}漏洞', style='MyHeading')
            doc.add_paragraph(f'漏洞类型:{vuln_type}', style='Normal_2').runs[0].bold = True
            
            if vuln_type in vuln_descriptions:
                vuln_description = vuln_descriptions[vuln_type]
                # 直接在段落级别设置首行缩进
                p = doc.add_paragraph(style='Normal')
                p.add_run(f'{vuln_description}')
                p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.line_spacing = 1.5  

            doc.add_paragraph(f'目标系统:{target_system}', style='Normal').runs[0].bold = True
            doc.add_paragraph(f'漏洞URL:{vuln_url}', style='Normal').runs[0].bold = True
            doc.add_paragraph(f'IP地址:{ip_address}', style='Normal').runs[0].bold = True

            # 添加资产证明信息
            if image_path_asset:
                try:
                    with Image.open(image_path_asset) as img:
                        width, height = img.size
                    doc.add_paragraph("1:资产归属证明", style='Normal')
                    doc.add_picture(image_path_asset, width=Inches(6), height=Inches(3))
                except Exception as img_error:
                    print(f"Error loading asset image: {img_error}")

            doc.add_paragraph(f'2:{vuln_proof}', style='Normal')

            # 添加漏洞证明信息
            if image_path_vuln:
                try:
                    with Image.open(image_path_vuln) as img:
                        width, height = img.size
                    doc.add_picture(image_path_vuln, width=Inches(6), height=Inches(3))
                except Exception as img_error:
                    print(f"Error loading vuln image: {img_error}")
            
            doc.add_paragraph(f'3:{vuln_proof_2}', style='Normal')
                    
            if image_path_vuln_2:
                try:
                    with Image.open(image_path_vuln_2) as img:
                        width, height = img.size
                    doc.add_picture(image_path_vuln_2, width=Inches(6), height=Inches(3))
                except Exception as img_error:
                    print(f"Error loading vuln image 2: {img_error}")
                    
            doc.add_paragraph(f'修复建议:{repair_suggestion}', style='Normal').runs[0].bold = True 

            report_name = f'{vuln_belongs}{vuln_type}漏洞情报线索.docx'
            doc.save(report_name)

            QMessageBox.information(self, '报告生成', f'漏洞报告已生成：{report_name}')

        except Exception as e:
            print(f"Error creating document: {e}")
            QMessageBox.warning(self, '错误', '无法创建文档')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon('th.jpg'))
    ex = ReportGenerator()
    ex.show()
    sys.exit(app.exec_())
